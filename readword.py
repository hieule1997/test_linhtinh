# import docx2txt

# text = docx2txt.process("READING-Part4.docx")

# #Prints output after converting

# content = []
# for line in text.splitlines():
#   #This will ignore empty/blank li 	nes. 
#   if line != '':
#     #Append to list
#     content.append(line)

# print (content)

#  Dua ra du lieu cac bang
from docx import Document

document = Document('/home/hieu/Desktop/cautrucde/Data/LISTENNING/LISTENING-Part2.docx')
tables = document.tables

for table in tables:
	fullquestions = ''
	cellquestions = table.rows[0] .cells[0]
	for paragraph in cellquestions.paragraphs:
		fullquestions = fullquestions + paragraph.text +"\n"
	print(fullquestions + "\n")
	rowsanser = table.rows[1:]
	i = 1
	for row in rowsanser:
		cellanser = row.cells[1:]

		for cell in cellanser:
			if i==1:
				for paragraph in cell.paragraphs:
					print("cau hoi")
					print(paragraph.text)
			elif i==2:
				for paragraph in cell.paragraphs:
					print("dung")
					print(paragraph.text)
			else:
				for paragraph in cell.paragraphs:
					print("sai")
					print(paragraph.text)
			if i == 4:
				i = 0
			i += 1 








# from subprocess import Popen, PIPE
# # from docx import opendocx, getdocumenttext
# # from cStringIO import StringIO
# def document_to_text(filename, file_path):
#     cmd = ['antiword', file_path]
#     p = Popen(cmd, stdout=PIPE)
#     stdout, stderr = p.communicate()
#     return stdout.decode('ascii', 'ignore')

# print (document_to_text('READING-Part3.doc','demo1.docx'))